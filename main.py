# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.
import os, glob
from tkinter import *
from tkinter import ttk
import fitz
import docx
from tkinter.messagebox import showerror, showwarning, showinfo


root = Tk()
root.geometry("600x400")

root.title('Search box - KA')
root.iconphoto(False,PhotoImage(file='support.png'))
root.columnconfigure(0,weight=2)
root.columnconfigure(1,weight=1)
root.columnconfigure(2,weight=1)
root.columnconfigure(3,weight=1)
#root.columnconfigure(1,weight=2)

# Add the textfield for folderpath
en = Entry(root, width=60, borderwidth=5)
en.insert(0, 'Enter the path of the folder')
en.grid(column=1,row=0,columnspan=3,padx=10,pady=20)
# Add the label for folderpath
label1 = Label(root,text='Nhập địa chỉ folder:')
label1.grid(column=0,row=0,padx=10,pady=20)


e = IntVar()
e.set(1)


# pdf radiobutton
pdfbutton = ttk.Radiobutton(root,text='pdf',value=103,variable=e)
pdfbutton.grid(column=1, row=2, ipadx=10, ipady=10)
# docx radiobutton
docxbutton = ttk.Radiobutton(root,text='docx',value=102,variable=e)
docxbutton.grid(column=2, row=2, ipadx=10, ipady=10)
# text radio button
txtbutton = ttk.Radiobutton(root,text='txt',value=101,variable=e)
txtbutton.grid(column=3, row=2, ipadx=10, ipady=10)
# label2 type of file
label2 = Label(root,text='Loại file cần tìm: ')
label2.grid(column=0,row=2,padx=5,pady=20)

# radiobutton số lượng từ tìm kiếm

labelw = Label(root,text="Số từ cần tìm: ")
labelw.grid(column=0,row=3,padx=30,pady=20,sticky=W)
many = IntVar()
many.set(20)


def deactivate():
    if many.get() == 1:
        sword2.configure(state='disabled')
        num2.configure(state='disabled')
    if many.get() == 2:
        sword2.configure(state='normal')
        num2.configure(state='normal')


one = ttk.Radiobutton(root,text='1',variable=many,value=1,command=deactivate)
one.grid(column=1,row=3,padx=20)

two = ttk.Radiobutton(root,text='2',variable=many,value=2,command=deactivate)
two.grid(column=2,row=3,padx=20)


# create frame to contain seeking words
frame1 = ttk.Frame(root)
#frame1['padding'] = (5,10)
#frame1['borderwidth'] = 5
frame1.grid(column=0,columnspan=4, row=4,pady=10,padx=20,sticky=W)
frame1.columnconfigure(0, weight=1)
frame1.columnconfigure(1, weight=3)
frame1.columnconfigure(2, weight=1)
frame1.columnconfigure(3, weight=1)

# searching box 1
label3 = Label(frame1,text="Cụm từ/từ thứ nhất:")
label3.grid(column=0,row=0,sticky=W)

label4 = Label(frame1,text="Số lần xuất hiện:")
label4.grid(column=2,row=0)

sword1= Entry(frame1,width=30)
sword1.grid(column=1,row=0)

num1 = Entry(frame1,width=10)
num1.grid(column=3,row=0)

# searchingbox 2
label5 = Label(frame1,text="Cụm từ/từ thứ hai:")
label5.grid(column=0,row=1,pady=20,sticky=W)

label6 = Label(frame1,text="Số lần xuất hiện:")
label6.grid(column=2,row=1,sticky=E)

sword2= Entry(frame1,width=30)
sword2.grid(column=1,row=1,padx=20)

num2 = Entry(frame1,width=10)
num2.grid(column=3,row=1,sticky=E)


# the main function
def required():
    state=False
    tail = ''
    fol = en.get()
    folPath = os.path.normpath(fol)
    if e.get() == 101:
        tail = 'txt'
    if e.get() == 102:
        tail = 'docx'
    if e.get() == 103:
        tail = 'pdf'
    # get all the filepath inside a folder, including files of subfolders and sub-subfolders
    fileg = glob.glob('{0}/**/*.{1}'.format(folPath, tail), recursive=True)

    if many.get() == 1:
        if sword1.get() == '' or num1.get() == '':
            state=False
            showerror('Error','Vui lòng điền vào các trường còn trống.')
        else:
            try:
                print('kimanh')
                cl = sword1.get()
                nu = int(num1.get())
                state=True
                if tail == 'docx':
                    protected = []
                    track = []
                    for file in fileg:
                        count = 0
                        try:
                            doc = docx.Document(file)
                            for para in doc.paragraphs:
                                para = para.text
                                count += count + para.count(cl) + para.count(cl.capitalize())
                            if count >= nu:
                                track.append(file)
                        except:
                            protected.append(file)
                if tail == 'txt':
                    track = []
                    protected = []
                    for file in fileg:
                        count = 0
                        try:
                            with open(file) as fhandle:
                                for line in fhandle:
                                    line = line.strip()
                                    if cl in line or cl.capitalize() in line:
                                        count += 1
                            if count >= nu:
                                track.append(file)
                        except:
                            protected.append(file)
                if tail=='pdf':
                    track=[]
                    protected=[]
                    for file in fileg:
                        try:
                            count=0
                            doc = fitz.open(file)
                            if len(doc) <= 10:
                                for page in doc:
                                    text = page.get_text('text')
                                    count = count + text.count(cl) + text.count(cl.capitalize())
                                if count >= nu:
                                    track.append(file)
                            else:
                                pnum = 0
                                for page in doc:
                                    while pnum <= 10:
                                        text = page.get_text('text')
                                        count = count+text.count(cl)+text.count(cl.capitalize())
                                        pnum += 1
                                if count >= nu:
                                    track.append(file)
                        except:
                            protected.append(file)
            except:
                state=False
                showerror('Error','Mục 2 và 4 chỉ nhận giá trị số.')

    else:
        if sword1.get() == '' or sword2.get() == '' or num1.get() == '' or num2.get() == '':
            state=False
            showerror('Error', 'Vui lòng điền vào các trường còn trống.')
        else:
            try:

                cl1 = sword1.get()
                nu1 = int(num1.get())
                cl2 = sword2.get()
                nu2 = int(num2.get())
                state=True
                if tail == 'docx':
                    protected = []
                    track = []
                    for file in fileg:
                        count1 = 0
                        count2 = 0
                        try:
                            doc = docx.Document(file)
                            for para in doc.paragraphs:
                                para = para.text
                                count1 += count1 + para.count(cl1) + para.count(cl1.capitalize())
                                count2 += count2 + para.count(cl2) + para.count(cl2.capitalize())
                            if count1 >= nu1 and count2 >= nu2:
                                track.append(file)
                        except:
                            protected.append(file)
                if tail == 'txt':
                    track = []
                    protected = []
                    for file in fileg:
                        count1 = 0
                        count2 = 0
                        try:
                            with open(file) as fhandle:
                                for line in fhandle:
                                    line = line.strip()
                                    if cl1 in line or cl1.capitalize() in line:
                                        count1 += 1
                                    if cl2 in line or cl2.capitalize() in line:
                                        count2 += 1
                            if count1 >= nu1 and count2 >= nu2:
                                track.append(file)
                        except:
                            protected.append(file)
                if tail == 'pdf':
                    track = []
                    protected = []
                    for file in fileg:
                        try:
                            count1 = 0
                            count2 = 0
                            doc = fitz.open(file)
                            if len(doc) <= 10:
                                for page in doc:
                                    text = page.get_text('text')
                                    count1 = count1 + text.count(cl1) + text.count(cl1.capitalize())
                                    count2 = count2 + text.count(cl2) + text.count(cl2.capitalize())
                                if count1 >= nu1 and count2 >= nu2:
                                    track.append(file)
                            else:
                                pnum = 0
                                for page in doc:
                                    while pnum <= 10:
                                        text = page.get_text('text')
                                        count1 = count1+text.count(cl1)+text.count(cl1.capitalize())
                                        count2 = count2 + text.count(cl2) + text.count(cl2.capitalize())
                                        pnum += 1
                                if count1 >= nu1 and count2 >= nu2:
                                    track.append(file)
                        except:
                            protected.append(file)
            except:
                state=False
                showerror('Error','Mục 2 và 4 chỉ nhận giá trị số.')



    # create a listbox to show the search result. The listbox will be in another tab
    # create a listbox to show the search result. The listbox will be in another tab
    if state == True:
        canvas = Tk()
        canvas.geometry('600x550')
        canvas.title('Search result')

        # my_listbox=Listbox(canvas,width=50)
        # my_listbox.pack()
        my_frame = Frame(canvas)
        my_frame.pack()

        refile = Label(my_frame,text="Các files thỏa mãn yêu cầu: ")
        refile.pack(pady=20,padx=10)
        my_scroll_bar = Scrollbar(my_frame, orient=VERTICAL)
        my_scroll_barh = Scrollbar(my_frame, orient=HORIZONTAL)
        my_listbox = Listbox(my_frame, width=80,height=10, yscrollcommand=my_scroll_bar.set,xscrollcommand=my_scroll_barh.set)

        my_scroll_bar.config(command=my_listbox.yview)
        my_scroll_bar.pack(side=RIGHT,fill=Y)

        my_scroll_barh.config(command=my_listbox.xview)
        my_scroll_barh.pack(side=BOTTOM, fill=X)

        my_listbox.pack()
        # make the track appear in the listbox
        for item in track:
            my_listbox.insert(END, item)

        def open_file():
            file = my_listbox.get(ANCHOR)
            os.startfile(file)

        # create button to open required files
        select_button = Button(canvas, text="Open file", command=open_file)
        select_button.pack(pady=15)
    # create another label
        plabel=Label(canvas,text='Các file được bảo vệ hoặc lỗi format: ')
        plabel.pack(pady=10)

        # create another listbox to display protected files
        my_frame2 = Frame(canvas)

        my_frame2.pack()
        my_scroll_bar2 = Scrollbar(my_frame2, orient=VERTICAL)
        my_scroll_bar2h = Scrollbar(my_frame2, orient=HORIZONTAL)
        my_listbox2 = Listbox(my_frame2, width=80,height=8, yscrollcommand=my_scroll_bar2.set,xscrollcommand=my_scroll_bar2h.set)
        my_scroll_bar2.config(command=my_listbox2.yview)
        my_scroll_bar2.pack(side=RIGHT, fill=Y)
        my_scroll_bar2h.config(command=my_listbox2.xview)
        my_scroll_bar2h.pack(side=BOTTOM, fill=X)
        my_listbox2.pack()

        for item in protected:
            my_listbox2.insert(END, item)

        def open_file2():
            file = my_listbox2.get(ANCHOR)
            os.startfile(file)

        select_button2 = Button(canvas, text="Open protected file", command=open_file2)
        select_button2.pack(pady=15)


buttonsubmit = Button(root, text='Submit',padx=10,pady=10, command=required)
buttonsubmit.grid(column=1,row=6,padx=60,sticky=E)

root.mainloop()
